from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title Page
title = doc.add_heading('Dynamic Question Generation System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
subtitle = doc.add_paragraph('Using RAG (Retrieval-Augmented Generation) and LLM')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

# Project Details
details = doc.add_paragraph()
details.add_run('A Project Report\n').bold = True
details.add_run('Submitted in partial fulfillment of the requirements\n')
details.add_run('for the degree of\n')
details.add_run('Bachelor of Engineering\n').bold = True
details.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# Table of Contents
doc.add_heading('TABLE OF CONTENTS', level=1)
toc_items = [
    ('1. Introduction', '3'),
    ('2. Problem Statement', '4'),
    ('3. Literature Review', '5'),
    ('4. System Requirements', '6'),
    ('5. System Architecture', '7'),
    ('6. Implementation', '8'),
    ('7. Results', '10'),
    ('8. Conclusion', '11'),
    ('9. Future Work', '12'),
    ('10. References', '13'),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.add_run(item)
    p.add_run('\t' * 8 + page)

doc.add_page_break()

# Chapter 1: Introduction
doc.add_heading('CHAPTER 1: INTRODUCTION', level=1)
doc.add_paragraph(
    'The Dynamic Question Generation System is an innovative educational technology solution '
    'that leverages Retrieval-Augmented Generation (RAG) and Large Language Models (LLMs) to '
    'automatically generate examination questions. This system addresses the critical need for '
    'efficient and diverse question generation in educational institutions.'
)
doc.add_paragraph(
    'Traditional methods of question paper preparation are time-consuming and often result in '
    'repetitive patterns that students can predict. Our system solves this by dynamically generating '
    'unique questions based on course content, difficulty levels, and topic specifications.'
)
doc.add_heading('1.1 Background', level=2)
doc.add_paragraph(
    'With the advancement of Natural Language Processing (NLP) and the emergence of powerful '
    'Large Language Models, automated question generation has become increasingly viable. '
    'RAG combines the benefits of information retrieval with generative AI, allowing the system '
    'to ground question generation in actual course material stored in a vector database.'
)

doc.add_heading('1.2 Objectives', level=2)
doc.add_paragraph('The main objectives of this project are:', style='List Bullet')
doc.add_paragraph('To develop a system that generates contextually relevant exam questions', style='List Bullet')
doc.add_paragraph('To implement difficulty-based question generation (Easy, Medium, Hard)', style='List Bullet')
doc.add_paragraph('To create a web-based interface for examiners and students', style='List Bullet')
doc.add_paragraph('To integrate local LLM (Ollama) for privacy-focused question generation', style='List Bullet')

doc.add_page_break()

# Chapter 2: Problem Statement
doc.add_heading('CHAPTER 2: PROBLEM STATEMENT', level=1)
doc.add_paragraph(
    'Educational institutions face several challenges in examination management:'
)
doc.add_paragraph('Manual question paper preparation is extremely time-consuming', style='List Bullet')
doc.add_paragraph('Limited variety leads to predictable patterns in examinations', style='List Bullet')
doc.add_paragraph('Difficulty in maintaining consistent difficulty levels across questions', style='List Bullet')
doc.add_paragraph('Inability to generate questions dynamically during examinations', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'This project aims to solve these problems by creating an automated, intelligent question '
    'generation system that can produce unique, topic-specific questions on demand while '
    'ensuring appropriate difficulty levels and content relevance.'
)

doc.add_page_break()

# Chapter 3: Literature Review
doc.add_heading('CHAPTER 3: LITERATURE REVIEW', level=1)

doc.add_heading('3.1 Natural Language Processing in Education', level=2)
doc.add_paragraph(
    'NLP has revolutionized educational technology, enabling automated essay grading, '
    'intelligent tutoring systems, and question generation. Recent advances in transformer '
    'architectures have significantly improved the quality of generated text.'
)

doc.add_heading('3.2 Retrieval-Augmented Generation (RAG)', level=2)
doc.add_paragraph(
    'RAG is a technique that combines information retrieval with text generation. It retrieves '
    'relevant context from a knowledge base and uses it to generate more accurate and grounded '
    'responses. This approach reduces hallucination in LLM outputs.'
)

doc.add_heading('3.3 Vector Databases and Embeddings', level=2)
doc.add_paragraph(
    'Vector databases like FAISS enable efficient similarity search over high-dimensional embeddings. '
    'Sentence transformers convert text into dense vectors that capture semantic meaning, '
    'allowing for contextual retrieval of relevant content.'
)

doc.add_page_break()

# Chapter 4: System Requirements
doc.add_heading('CHAPTER 4: SYSTEM REQUIREMENTS', level=1)

doc.add_heading('4.1 Software Requirements', level=2)
doc.add_paragraph('Python 3.10 or higher', style='List Bullet')
doc.add_paragraph('Flask Web Framework', style='List Bullet')
doc.add_paragraph('FAISS Vector Database', style='List Bullet')
doc.add_paragraph('Sentence Transformers (all-MiniLM-L6-v2)', style='List Bullet')
doc.add_paragraph('Ollama with LLaMA 3 Model', style='List Bullet')
doc.add_paragraph('HTML/CSS/JavaScript for Frontend', style='List Bullet')

doc.add_heading('4.2 Hardware Requirements', level=2)
doc.add_paragraph('Processor: Intel Core i5 or equivalent', style='List Bullet')
doc.add_paragraph('RAM: 16 GB minimum (for LLM inference)', style='List Bullet')
doc.add_paragraph('Storage: 10 GB for model files', style='List Bullet')
doc.add_paragraph('GPU: Optional (improves inference speed)', style='List Bullet')

doc.add_page_break()

# Chapter 5: System Architecture
doc.add_heading('CHAPTER 5: SYSTEM ARCHITECTURE', level=1)

doc.add_heading('5.1 High-Level Architecture', level=2)
doc.add_paragraph(
    'The system follows a modular architecture with the following components:'
)
doc.add_paragraph('Web Interface Layer (Flask + HTML/CSS)', style='List Bullet')
doc.add_paragraph('Application Logic Layer (Question Generator Service)', style='List Bullet')
doc.add_paragraph('RAG Pipeline (Retriever + Prompt Builder)', style='List Bullet')
doc.add_paragraph('LLM Integration Layer (Ollama API)', style='List Bullet')
doc.add_paragraph('Data Layer (FAISS Index + Metadata)', style='List Bullet')

doc.add_heading('5.2 Data Flow', level=2)
doc.add_paragraph(
    '1. User selects subject, topics, and difficulty level\n'
    '2. System retrieves relevant context from vector database\n'
    '3. Prompt is constructed with context and rules\n'
    '4. LLM generates question based on prompt\n'
    '5. Question is displayed to the user'
)

doc.add_page_break()

# Chapter 6: Implementation
doc.add_heading('CHAPTER 6: IMPLEMENTATION', level=1)

doc.add_heading('6.1 Text Extraction and Preprocessing', level=2)
doc.add_paragraph(
    'Course materials are extracted from PDF documents and cleaned to remove formatting artifacts. '
    'The text is then chunked into manageable segments for embedding.'
)

doc.add_heading('6.2 Vector Database Creation', level=2)
doc.add_paragraph(
    'Text chunks are embedded using Sentence Transformers and stored in a FAISS index. '
    'Metadata including subject, topic, and source is preserved for filtered retrieval.'
)

doc.add_heading('6.3 RAG Retriever', level=2)
doc.add_paragraph(
    'The retriever module queries the FAISS index with the selected subject and topic, '
    'returning the most relevant context chunks for question generation.'
)

doc.add_heading('6.4 Prompt Builder', level=2)
doc.add_paragraph(
    'The prompt builder constructs a structured prompt containing the retrieved context, '
    'difficulty rules, and generation instructions for the LLM.'
)

doc.add_heading('6.5 LLM Question Generator', level=2)
doc.add_paragraph(
    'The question generator module communicates with the Ollama API to generate questions. '
    'It handles connection errors and model availability gracefully.'
)

doc.add_heading('6.6 Web Interface', level=2)
doc.add_paragraph(
    'A Flask-based web interface provides separate views for examiners (to configure exams) '
    'and students (to take tests). Session management tracks student progress.'
)

doc.add_page_break()

# Chapter 7: Results
doc.add_heading('CHAPTER 7: RESULTS', level=1)
doc.add_paragraph(
    'The system successfully generates diverse, contextually relevant questions based on '
    'the input parameters. Key achievements include:'
)
doc.add_paragraph('Dynamic question generation from stored course content', style='List Bullet')
doc.add_paragraph('Configurable difficulty levels affecting question complexity', style='List Bullet')
doc.add_paragraph('Real-time question generation during examinations', style='List Bullet')
doc.add_paragraph('Clean web interface for both examiners and students', style='List Bullet')
doc.add_paragraph('Navigation between questions with history tracking', style='List Bullet')

doc.add_page_break()

# Chapter 8: Conclusion
doc.add_heading('CHAPTER 8: CONCLUSION', level=1)
doc.add_paragraph(
    'This project successfully demonstrates the application of RAG and LLM technologies '
    'for automated question generation in educational settings. The system provides a '
    'practical solution to the challenges of manual question paper preparation while '
    'ensuring variety, relevance, and appropriate difficulty in generated questions.'
)
doc.add_paragraph(
    'The use of local LLM through Ollama ensures data privacy and reduces dependency on '
    'external APIs. The modular architecture allows for easy extension and customization.'
)

doc.add_page_break()

# Chapter 9: Future Work
doc.add_heading('CHAPTER 9: FUTURE WORK', level=1)
doc.add_paragraph('Support for multiple question types (MCQ, True/False, Fill-in-the-blanks)', style='List Bullet')
doc.add_paragraph('Answer key generation for generated questions', style='List Bullet')
doc.add_paragraph('Automatic grading of student responses', style='List Bullet')
doc.add_paragraph('Integration with Learning Management Systems (LMS)', style='List Bullet')
doc.add_paragraph('Support for image-based questions', style='List Bullet')
doc.add_paragraph('Analytics dashboard for question performance', style='List Bullet')

doc.add_page_break()

# Chapter 10: References
doc.add_heading('CHAPTER 10: REFERENCES', level=1)
doc.add_paragraph('[1] Lewis, P., et al. "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks." NeurIPS 2020.')
doc.add_paragraph('[2] Vaswani, A., et al. "Attention Is All You Need." NeurIPS 2017.')
doc.add_paragraph('[3] Reimers, N., & Gurevych, I. "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks." EMNLP 2019.')
doc.add_paragraph('[4] Johnson, J., Douze, M., & Jégou, H. "Billion-scale similarity search with GPUs." IEEE Transactions on Big Data, 2019.')
doc.add_paragraph('[5] Flask Documentation. https://flask.palletsprojects.com/')
doc.add_paragraph('[6] Ollama Documentation. https://ollama.com/')

# Save the document
doc.save(r'd:\Dynamic-Question-generation-RAG\Project_Report.docx')
print("Report saved successfully!")
